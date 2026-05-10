import base64
from html import escape as _he
from pathlib import Path
from urllib.parse import quote

import requests as _req

from PyQt6.QtCore import Qt, QSizeF, QThread, QTimer, QUrl, pyqtSignal
from PyQt6.QtGui import (
    QColor, QDesktopServices, QGuiApplication, QImage, QPainter, QPixmap, QTextDocument,
)
from PyQt6.QtWidgets import (
    QApplication, QCheckBox, QFileDialog, QFrame, QHBoxLayout, QLabel, QLineEdit,
    QMenu, QMessageBox, QPushButton, QScrollArea, QSizePolicy, QTextEdit,
    QVBoxLayout, QWidget,
)

import keep4mac.i18n as i18n
from keep4mac.api.keep_client import KeepClient
from keep4mac.core.models import COLOR_HEX, NoteColor, NoteModel, NoteType
from keep4mac.core.url_utils import extract_urls, short_url
from keep4mac.i18n import gettext as _


class _IMELineEdit(QLineEdit):
    """macOS 한국어 IME 첫 글자 자소 분리 방지 — focusIn 시 IME 재초기화."""
    def focusInEvent(self, event):
        super().focusInEvent(event)
        QGuiApplication.inputMethod().reset()


class _IMETextEdit(QTextEdit):
    """macOS 한국어 IME 첫 글자 자소 분리 방지 — focusIn 시 IME 재초기화."""
    def focusInEvent(self, event):
        super().focusInEvent(event)
        QGuiApplication.inputMethod().reset()


class _TranslateThread(QThread):
    """MyMemory 무료 API를 사용해 제목·본문을 백그라운드 번역한다."""
    done = pyqtSignal(str, str)   # (translated_title, translated_content)
    error = pyqtSignal(str)

    _API = "https://api.mymemory.translated.net/get"

    def __init__(self, title: str, content: str, source: str, target: str):
        super().__init__()
        self._title = title
        self._content = content
        self._source = source
        self._target = target

    def _call(self, text: str) -> str:
        if not text.strip():
            return text
        resp = _req.get(
            self._API,
            params={"q": text[:2000], "langpair": f"{self._source}|{self._target}"},
            timeout=12,
        )
        resp.raise_for_status()
        data = resp.json()
        if data.get("responseStatus") == 200:
            return data["responseData"]["translatedText"]
        raise RuntimeError(data.get("responseDetails", "Translation error"))

    def run(self):
        try:
            self.done.emit(self._call(self._title), self._call(self._content))
        except Exception as e:
            self.error.emit(str(e))


class _CreateNoteThread(QThread):
    """번역된 노트를 백그라운드에서 Google Keep에 생성하고 동기화한다."""
    done = pyqtSignal()
    error = pyqtSignal(str)

    def __init__(self, client, title: str, content: str, color):
        super().__init__()
        self._client = client
        self._title = title
        self._content = content
        self._color = color

    def run(self):
        try:
            self._client.create_note(self._title, self._content, self._color)
            self._client.sync()
            self.done.emit()
        except Exception as e:
            self.error.emit(str(e))


def _write_hwpx(path: str, title: str, body: str,
                checklist: "list[tuple[bool, str]] | None" = None,
                img_data: bytes | None = None) -> None:
    """python-hwpx(airmang) 라이브러리로 .hwpx 파일을 생성한다."""
    import io, contextlib
    from hwpx import HwpxDocument

    # 라이브러리가 Skeleton 로딩 시 출력하는 fallback 경고 억제
    with contextlib.redirect_stderr(io.StringIO()):
        doc = HwpxDocument.new()

    # 스켈레톤에 기본 빈 단락이 하나 있으므로 텍스트로 덮어쓰기
    first = True

    def _add(text: str, style_id: int = 0):
        nonlocal first
        if first:
            p = doc.sections[0].paragraphs[0]
            p.clear_text()
            p.add_run(text)
            if style_id:
                p.style_id_ref = str(style_id)
            first = False
            return p
        return doc.add_paragraph(text, style_id_ref=style_id if style_id else None)

    if title:
        _add(title, style_id=2)  # 개요 1 스타일 (큰 제목)

    if checklist:
        for is_chk, txt in checklist:
            _add(f"{'☑' if is_chk else '☐'} {txt}")
    elif body:
        for line in (body.splitlines() or [""]):
            _add(line)

    doc.save_to_path(path)

    if img_data:
        _inject_hwpx_image(path, img_data)


def _inject_hwpx_image(path: str, img_data: bytes) -> None:
    """python-hwpx가 생성한 HWPX ZIP에 이미지 단락을 직접 주입한다."""
    import os, random, zipfile
    from PyQt6.QtGui import QImage

    qi = QImage()
    qi.loadFromData(img_data)
    w_px, h_px = max(1, qi.width()), max(1, qi.height())
    org_w, org_h = w_px * 75, h_px * 75   # 1px = 75 HWP unit (96dpi)
    MAX_W = 42520                           # A4 콘텐츠 폭
    if org_w > MAX_W:
        cur_h = int(org_h * MAX_W / org_w)
        cur_w = MAX_W
    else:
        cur_w, cur_h = org_w, org_h

    sc_x = round(cur_w / org_w, 6)
    sc_y = round(cur_h / org_h, 6)
    cx, cy = cur_w // 2, cur_h // 2
    instid = random.randint(10_000_000, 99_999_999)
    pic_id = random.randint(1_000_000_000, 2_000_000_000)
    bin_id = "image1"

    img_para = (
        f'<hp:p id="{pic_id}" paraPrIDRef="0" styleIDRef="0"'
        f' pageBreak="0" columnBreak="0" merged="0">'
        f'<hp:run charPrIDRef="0">'
        f'<hp:pic id="{pic_id}" zOrder="0" numberingType="PICTURE"'
        f' textWrap="SQUARE" textFlow="BOTH_SIDES" lock="0"'
        f' dropcapstyle="None" href="" groupLevel="0" instid="{instid}" reverse="0">'
        f'<hp:offset x="0" y="0"/>'
        f'<hp:orgSz width="{org_w}" height="{org_h}"/>'
        f'<hp:curSz width="{cur_w}" height="{cur_h}"/>'
        f'<hp:flip horizontal="0" vertical="0"/>'
        f'<hp:rotationInfo angle="0" centerX="{cx}" centerY="{cy}" rotateimage="1"/>'
        f'<hp:renderingInfo>'
        f'<hc:transMatrix e1="1" e2="0" e3="0" e4="0" e5="1" e6="0"/>'
        f'<hc:scaMatrix e1="{sc_x}" e2="0" e3="0" e4="0" e5="{sc_y}" e6="0"/>'
        f'<hc:rotMatrix e1="1" e2="0" e3="0" e4="0" e5="1" e6="0"/>'
        f'</hp:renderingInfo>'
        f'<hc:img binaryItemIDRef="{bin_id}" bright="0" contrast="0"'
        f' effect="REAL_PIC" alpha="0"/>'
        f'<hp:imgRect>'
        f'<hc:pt0 x="0" y="0"/>'
        f'<hc:pt1 x="{org_w}" y="0"/>'
        f'<hc:pt2 x="{org_w}" y="{org_h}"/>'
        f'<hc:pt3 x="0" y="{org_h}"/>'
        f'</hp:imgRect>'
        f'<hp:imgClip left="0" right="{org_w}" top="0" bottom="{org_h}"/>'
        f'<hp:inMargin left="0" right="0" top="0" bottom="0"/>'
        f'<hp:imgDim dimwidth="{org_w}" dimheight="{org_h}"/>'
        f'<hp:effects/>'
        f'<hp:sz width="{cur_w}" widthRelTo="ABSOLUTE"'
        f' height="{cur_h}" heightRelTo="ABSOLUTE" protect="0"/>'
        f'<hp:pos treatAsChar="0" affectLSpacing="0" flowWithText="1"'
        f' allowOverlap="0" holdAnchorAndSO="0"'
        f' vertRelTo="PARA" horzRelTo="PARA"'
        f' vertAlign="TOP" horzAlign="LEFT" vertOffset="0" horzOffset="0"/>'
        f'<hp:outMargin left="0" right="0" top="0" bottom="0"/>'
        f'<hp:shapeComment/>'
        f'</hp:pic>'
        f'</hp:run>'
        f'<hp:linesegarray>'
        f'<hp:lineseg textpos="0" vertpos="0" vertsize="1000" textheight="1000"'
        f' baseline="850" spacing="600" horzpos="0" horzsize="{MAX_W}" flags="393216"/>'
        f'</hp:linesegarray>'
        f'</hp:p>'
    )

    with zipfile.ZipFile(path, 'r') as zin:
        names = zin.namelist()
        files = {n: zin.read(n) for n in names}

    files[f'BinData/{bin_id}.png'] = img_data

    # 매니페스트에 이미지 항목 추가
    hpf = files['Contents/content.hpf'].decode('utf-8')
    item_xml = (
        f'<opf:item id="{bin_id}" href="BinData/{bin_id}.png"'
        f' media-type="image/png" isEmbeded="1"/>'
    )
    hpf = hpf.replace('</opf:manifest>', item_xml + '</opf:manifest>')
    files['Contents/content.hpf'] = hpf.encode('utf-8')

    # section0.xml에 hc: 네임스페이스 없으면 추가
    sec = files['Contents/section0.xml'].decode('utf-8')
    if 'xmlns:hc=' not in sec:
        HC = 'xmlns:hc="http://www.hancom.co.kr/hwpml/2011/core"'
        sec = sec.replace('<hs:sec ', f'<hs:sec {HC} ', 1)

    # 첫 번째 </hp:p> 뒤에 이미지 단락 삽입
    pos = sec.find('</hp:p>')
    if pos != -1:
        insert_at = pos + len('</hp:p>')
        sec = sec[:insert_at] + img_para + sec[insert_at:]
    files['Contents/section0.xml'] = sec.encode('utf-8')

    # ZIP 재기록 (mimetype은 첫 번째, 비압축)
    tmp = path + '.tmp'
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        mime_info = zipfile.ZipInfo('mimetype')
        mime_info.compress_type = zipfile.ZIP_STORED
        zout.writestr(mime_info, files.pop('mimetype', b'application/hwp+zip'))
        for name, data in files.items():
            zout.writestr(name, data)
    os.replace(tmp, path)


def _write_docx(path: str, title: str, body: str,
                checklist: "list[tuple[bool, str]] | None" = None,
                img_data: bytes | None = None) -> None:
    """python-docx로 .docx 파일을 생성한다. HWP 2014+에서 열 수 있다."""
    import io
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 페이지 여백 설정 (A4 기본)
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    if title:
        h = doc.add_heading(title, level=1)
        h.runs[0].font.size = Pt(18)

    if img_data:
        try:
            img_stream = io.BytesIO(img_data)
            doc.add_picture(img_stream, width=Cm(14))
            doc.add_paragraph()
        except Exception:
            pass

    if checklist:
        for is_chk, text in checklist:
            sym = "☑" if is_chk else "☐"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{sym} {text}")
    elif body:
        for line in body.splitlines():
            doc.add_paragraph(line)

    doc.save(path)


# 색상 팔레트 순서 (DEFAULT는 흰색 맨 앞)
_PALETTE = [
    NoteColor.DEFAULT, NoteColor.RED, NoteColor.ORANGE, NoteColor.YELLOW,
    NoteColor.GREEN, NoteColor.TEAL, NoteColor.BLUE, NoteColor.CERULEAN,
    NoteColor.PURPLE, NoteColor.PINK, NoteColor.BROWN, NoteColor.GRAY,
]


class _ImageThread(QThread):
    done = pyqtSignal(bytes)

    def __init__(self, url: str, client: KeepClient):
        super().__init__()
        self._url = url
        self._client = client

    def run(self):
        data = self._client.fetch_image(self._url)
        if data:
            self.done.emit(data)


class NoteEditorWidget(QWidget):
    back_requested = pyqtSignal()

    def __init__(self, client: KeepClient):
        super().__init__()
        self._client = client
        self._note_id: str | None = None
        self._is_new = False
        self._is_pinned = False
        self._image_url: str | None = None
        self._current_color = NoteColor.DEFAULT
        self._color_btns: dict[NoteColor, QPushButton] = {}
        self._checklist_rows: list[tuple[QCheckBox, str]] = []
        self._img_thread: _ImageThread | None = None
        self._translate_thread: _TranslateThread | None = None
        self._orig_title: str = ""
        self._orig_body: str = ""
        self._revert_btn: QPushButton | None = None
        self._build_ui()
        self._title_edit.textChanged.connect(self._update_revert_btn)
        self._body_edit.textChanged.connect(self._update_revert_btn)

    # ── UI 구성 ───────────────────────────────────────────────

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # 헤더
        header = QWidget()
        header.setFixedHeight(48)
        header.setStyleSheet("background: #f5f5f7; border-bottom: 1px solid #d1d1d6;")
        hl = QHBoxLayout(header)
        hl.setContentsMargins(8, 0, 8, 0)
        hl.setSpacing(4)

        back_btn = QPushButton("←")
        back_btn.setFixedSize(32, 32)
        back_btn.setStyleSheet("""
            QPushButton { background: transparent; color: #007AFF; border: none; font-size: 16px; }
            QPushButton:hover { background: rgba(0,0,0,0.06); border-radius: 16px; }
        """)
        back_btn.clicked.connect(self._on_back)
        hl.addWidget(back_btn)

        self._header_label = QLabel(_("Edit Note"))
        self._header_label.setStyleSheet(
            "font-size: 14px; font-weight: 600; color: #1c1c1e; background: transparent;"
        )
        hl.addWidget(self._header_label, 1)

        root.addWidget(header)

        # 본문
        self._body_widget = QWidget()
        self._body_widget.setStyleSheet("background: white;")
        body = self._body_widget
        bl = QVBoxLayout(body)
        bl.setContentsMargins(16, 12, 16, 12)
        bl.setSpacing(10)

        self._title_edit = _IMELineEdit()
        self._title_edit.setPlaceholderText(_("Title"))
        self._title_edit.setStyleSheet("""
            QLineEdit {
                font-size: 16px; font-weight: 600; color: #1c1c1e;
                border: none; border-bottom: 1px solid #d1d1d6;
                padding: 4px 0; background: transparent;
            }
        """)
        bl.addWidget(self._title_edit)

        # 텍스트 노트 편집기
        self._body_edit = _IMETextEdit()
        self._body_edit.setPlaceholderText(_("Enter note content…"))
        self._body_edit.setStyleSheet("""
            QTextEdit { font-size: 13px; color: #1c1c1e; border: none; background: transparent; }
        """)
        self._body_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        bl.addWidget(self._body_edit, 1)

        # 체크리스트 노트 스크롤 영역
        self._list_scroll = QScrollArea()
        self._list_scroll.setWidgetResizable(True)
        self._list_scroll.setFrameShape(QFrame.Shape.NoFrame)
        self._list_scroll.setStyleSheet("background: transparent;")
        self._list_container = QWidget()
        self._list_container.setStyleSheet("background: transparent;")
        self._list_layout = QVBoxLayout(self._list_container)
        self._list_layout.setContentsMargins(0, 0, 0, 0)
        self._list_layout.setSpacing(6)
        self._list_layout.addStretch()
        self._list_scroll.setWidget(self._list_container)
        self._list_scroll.hide()
        bl.addWidget(self._list_scroll, 1)

        # 이미지 섹션
        self._img_section = QFrame()
        self._img_section.setStyleSheet(
            "QFrame { background: #f5f5f7; border: 1px solid #d1d1d6; border-radius: 10px; }"
        )
        img_l = QVBoxLayout(self._img_section)
        img_l.setContentsMargins(8, 8, 8, 8)
        img_l.setSpacing(6)

        self._img_display = QLabel()
        self._img_display.setFixedHeight(140)
        self._img_display.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._img_display.setStyleSheet("background: #e5e5ea; border-radius: 8px;")
        img_l.addWidget(self._img_display)

        img_del_btn = QPushButton(_("🗑  Delete Image"))
        img_del_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #FF3B30;
                border: 1px solid #FF3B30; border-radius: 8px;
                font-size: 12px; padding: 4px 10px;
            }
            QPushButton:hover { background: rgba(255,59,48,0.08); }
        """)
        img_del_btn.clicked.connect(self._on_delete_image)
        img_l.addWidget(img_del_btn, 0, Qt.AlignmentFlag.AlignRight)

        self._img_section.hide()
        bl.addWidget(self._img_section)

        # 링크 섹션
        self._links_section = QFrame()
        self._links_section.setStyleSheet(
            "QFrame { background: #f5f5f7; border: 1px solid #d1d1d6; border-radius: 10px; }"
        )
        links_l = QVBoxLayout(self._links_section)
        links_l.setContentsMargins(10, 8, 10, 8)
        links_l.setSpacing(4)

        links_header = QLabel(_("🔗 Links"))
        links_header.setStyleSheet(
            "font-size: 11px; font-weight: 600; color: #636366; background: transparent;"
        )
        links_l.addWidget(links_header)

        self._links_body = QVBoxLayout()
        self._links_body.setSpacing(2)
        links_l.addLayout(self._links_body)

        self._links_section.hide()
        bl.addWidget(self._links_section)

        root.addWidget(body, 1)

        # 하단 footer (색상 팔레트 2행 + 저장 버튼)
        self._footer_widget = QWidget()
        footer = self._footer_widget
        footer.setFixedHeight(64)
        footer.setStyleSheet("background: #f5f5f7; border-top: 1px solid #d1d1d6;")
        fl = QHBoxLayout(footer)
        fl.setContentsMargins(12, 8, 12, 8)
        fl.setSpacing(8)

        # 색상 팔레트 6×2 그리드
        palette_col = QVBoxLayout()
        palette_col.setSpacing(4)
        row1, row2 = _PALETTE[:6], _PALETTE[6:]
        for row_colors in (row1, row2):
            row_layout = QHBoxLayout()
            row_layout.setSpacing(4)
            row_layout.setContentsMargins(0, 0, 0, 0)
            for color in row_colors:
                btn = QPushButton()
                btn.setFixedSize(20, 20)
                btn.setCheckable(True)
                btn.clicked.connect(lambda _, c=color: self._on_color_pick(c))
                self._color_btns[color] = btn
                row_layout.addWidget(btn)
            palette_col.addLayout(row_layout)

        fl.addLayout(palette_col)

        # 핀 고정 토글 버튼
        self._pin_btn = QPushButton("📌")
        self._pin_btn.setFixedSize(32, 32)
        self._pin_btn.setCheckable(True)
        self._pin_btn.setToolTip(_("Pin / Unpin"))
        self._pin_btn.clicked.connect(self._on_pin_toggle)
        fl.addWidget(self._pin_btn, 0, Qt.AlignmentFlag.AlignVCenter)

        self._delete_btn = QPushButton("✕")
        self._delete_btn.setFixedSize(32, 32)
        self._delete_btn.setToolTip(_("Delete"))
        self._delete_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #FF3B30;
                border: 1.5px solid #FF3B30; border-radius: 8px;
                font-size: 15px; font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,59,48,0.08); }
            QPushButton:pressed { background: rgba(255,59,48,0.16); }
        """)
        self._delete_btn.clicked.connect(self._on_delete)
        fl.addWidget(self._delete_btn, 0, Qt.AlignmentFlag.AlignVCenter)

        self._copy_btn = QPushButton("📋")
        self._copy_btn.setFixedSize(32, 32)
        self._copy_btn.setToolTip(_("Copy to clipboard"))
        self._copy_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #636366;
                border: 1px solid #d1d1d6; border-radius: 8px; font-size: 14px;
            }
            QPushButton:hover { background: #f2f2f7; }
            QPushButton:pressed { background: #e5e5ea; }
        """)
        self._copy_btn.clicked.connect(self._on_copy_to_clipboard)
        fl.addWidget(self._copy_btn, 0, Qt.AlignmentFlag.AlignVCenter)

        fl.addStretch()
        self._refresh_palette()
        self._refresh_pin_btn()

        self._export_btn = QPushButton("📤")
        self._export_btn.setFixedSize(32, 32)
        self._export_btn.setToolTip(_("Export / Share"))
        self._export_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #636366;
                border: 1px solid #d1d1d6; border-radius: 8px; font-size: 14px;
            }
            QPushButton:hover { background: #f2f2f7; }
            QPushButton:pressed { background: #e5e5ea; }
        """)
        self._export_btn.clicked.connect(self._on_export_click)
        fl.addWidget(self._export_btn, 0, Qt.AlignmentFlag.AlignVCenter)

        self._revert_btn = QPushButton("↩")
        self._revert_btn.setFixedSize(32, 32)
        self._revert_btn.setToolTip(_("Revert"))
        self._revert_btn.setVisible(False)
        self._revert_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #636366;
                border: 1px solid #d1d1d6; border-radius: 8px; font-size: 14px;
            }
            QPushButton:hover { background: #f2f2f7; }
            QPushButton:pressed { background: #e5e5ea; }
        """)
        self._revert_btn.clicked.connect(self._on_revert)
        fl.addWidget(self._revert_btn, 0, Qt.AlignmentFlag.AlignVCenter)

        root.addWidget(footer)

    # ── 공개 API ──────────────────────────────────────────────

    def retranslate_ui(self):
        self._title_edit.setPlaceholderText(_("Title"))
        self._body_edit.setPlaceholderText(_("Enter note content…"))
        self._delete_btn.setToolTip(_("Delete"))
        self._copy_btn.setToolTip(_("Copy to clipboard"))
        self._pin_btn.setToolTip(_("Pin / Unpin"))
        self._export_btn.setToolTip(_("Export / Share"))
        self._revert_btn.setToolTip(_("Revert"))
        if self._is_new:
            self._header_label.setText(_("New Note"))
        else:
            self._header_label.setText(_("Edit Note"))

    def load_note(self, note_id: str):
        note = self._client.get_note(note_id)
        if not note:
            self.back_requested.emit()
            return
        self._note_id = note_id
        self._is_new = False
        self._header_label.setText(_("Edit Note"))
        self._delete_btn.show()
        self._populate(note)
        self._orig_title, self._orig_body = self._note_content()
        self._update_revert_btn()

    def new_note(self):
        self._note_id = None
        self._is_new = True
        self._image_url = None
        self._is_pinned = False
        self._header_label.setText(_("New Note"))
        self._delete_btn.hide()
        self._title_edit.clear()
        self._body_edit.clear()
        self._body_edit.show()
        self._list_scroll.hide()
        self._checklist_rows.clear()
        self._img_section.hide()
        self._links_section.hide()
        self._set_color(NoteColor.DEFAULT)
        self._refresh_pin_btn()
        self._orig_title = ""
        self._orig_body = ""
        self._update_revert_btn()
        self._title_edit.setFocus()

    # ── 내부 ──────────────────────────────────────────────────

    def _populate(self, note: NoteModel):
        self._image_url = note.image_url
        self._is_pinned = note.pinned
        self._title_edit.setText(note.title)
        self._set_color(note.color)
        self._refresh_pin_btn()

        if note.note_type == NoteType.LIST:
            self._body_edit.hide()
            self._list_scroll.show()
            self._rebuild_checklist(note)
            all_text = " ".join(i.text for i in note.checklist_items)
        else:
            self._list_scroll.hide()
            self._body_edit.show()
            self._body_edit.setPlainText(note.text)
            all_text = note.text

        self._update_image(note.image_url)
        self._update_links(all_text)

    def _rebuild_checklist(self, note: NoteModel):
        while self._list_layout.count() > 1:
            item = self._list_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._checklist_rows.clear()

        for ci in note.checklist_items:
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            rl = QHBoxLayout(row)
            rl.setContentsMargins(0, 0, 0, 0)
            rl.setSpacing(8)

            cb = QCheckBox()
            cb.setChecked(ci.checked)
            rl.addWidget(cb)

            lbl = QLabel(ci.text)
            lbl.setStyleSheet("font-size: 13px; color: #1c1c1e; background: transparent;")
            lbl.setWordWrap(True)
            rl.addWidget(lbl, 1)

            self._list_layout.insertWidget(self._list_layout.count() - 1, row)
            self._checklist_rows.append((cb, ci.text))

    def _update_image(self, image_url: str | None):
        if not image_url:
            self._img_section.hide()
            return
        self._img_display.clear()
        self._img_display.setText(_("Loading image…"))
        self._img_section.show()

        self._img_thread = _ImageThread(image_url, self._client)
        self._img_thread.done.connect(self._on_image_ready)
        self._img_thread.start()

    def _on_image_ready(self, data: bytes):
        pixmap = QPixmap()
        pixmap.loadFromData(data)
        if pixmap.isNull():
            return
        w = self._img_display.width() or 280
        h = self._img_display.height()
        scaled = pixmap.scaled(w, h, Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                               Qt.TransformationMode.SmoothTransformation)
        self._img_display.setText("")
        self._img_display.setPixmap(scaled)

    def _update_links(self, text: str):
        # 기존 링크 위젯 제거
        while self._links_body.count():
            item = self._links_body.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        urls = extract_urls(text)
        if not urls:
            self._links_section.hide()
            return

        for url in urls:
            display = short_url(url, 45)
            lbl = QLabel(f'<a href="{url}" style="color:#007AFF;">{display}</a>')
            lbl.setOpenExternalLinks(True)
            lbl.setWordWrap(True)
            lbl.setStyleSheet("font-size: 12px; background: transparent;")
            self._links_body.addWidget(lbl)

        self._links_section.show()

    # ── 슬롯 ──────────────────────────────────────────────────

    def _has_unsaved_changes(self) -> bool:
        title, body = self._note_content()
        return title != self._orig_title or body != self._orig_body

    def _update_revert_btn(self, *_):
        if self._revert_btn:
            self._revert_btn.setVisible(self._has_unsaved_changes())

    def _on_revert(self):
        self._title_edit.setText(self._orig_title)
        if self._list_scroll.isVisible() and self._note_id:
            note = self._client.get_note(self._note_id)
            if note:
                self._rebuild_checklist(note)
        else:
            self._body_edit.setPlainText(self._orig_body)
        self._update_revert_btn()

    def _do_save(self) -> bool:
        """저장만 수행한다. back_requested는 emit하지 않는다."""
        QGuiApplication.inputMethod().commit()
        title = self._title_edit.text().strip()
        body = self._body_edit.toPlainText()
        color = self._current_color

        if self._is_new:
            if not title and not body.strip():
                return False
            note = self._client.create_note(title, body, color, self._is_pinned)
            self._is_new = False
            self._note_id = note.id
            self._delete_btn.show()
        elif self._note_id:
            if self._list_scroll.isVisible():
                items = [(text, cb.isChecked()) for cb, text in self._checklist_rows]
                self._client.update_checklist(self._note_id, title, items, color)
            else:
                self._client.update_note(self._note_id, title, body, color)

        self._orig_title, self._orig_body = self._note_content()
        self._update_revert_btn()
        return True

    def auto_save_if_needed(self):
        """패널이 재표시되기 직전 호출 — 편집 중인 내용을 자동 저장한다."""
        if self._has_unsaved_changes():
            self._do_save()

    def _on_back(self):
        if self._has_unsaved_changes():
            self._do_save()
        self.back_requested.emit()

    def _on_pin_toggle(self):
        if self._is_new:
            # 새 노트는 저장 시 반영 — 버튼 상태만 업데이트
            self._is_pinned = not self._is_pinned
        else:
            # 기존 노트는 즉시 저장
            self._is_pinned = self._client.toggle_pin(self._note_id)
        self._refresh_pin_btn()

    def _refresh_pin_btn(self):
        self._pin_btn.setChecked(self._is_pinned)
        if self._is_pinned:
            self._pin_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(0,122,255,0.10); border: 1.5px solid #007AFF;
                    border-radius: 8px; font-size: 15px;
                }
                QPushButton:hover { background: rgba(0,122,255,0.18); }
            """)
        else:
            self._pin_btn.setStyleSheet("""
                QPushButton {
                    background: transparent; border: 1px solid #d1d1d6;
                    border-radius: 8px; font-size: 15px; color: #8e8e93;
                }
                QPushButton:hover { background: #f2f2f7; border-color: #8e8e93; }
            """)

    def _on_color_pick(self, color: NoteColor):
        self._set_color(color)

    def _set_color(self, color: NoteColor):
        self._current_color = color
        self._refresh_palette()
        self._apply_bg(color)

    def _refresh_palette(self):
        for color, btn in self._color_btns.items():
            hex_color = COLOR_HEX[color]
            is_selected = (color == self._current_color)
            check = "✓" if is_selected else ""
            border = "2px solid #1c1c1e" if is_selected else "1px solid rgba(0,0,0,0.18)"
            btn.setText(check)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {hex_color};
                    border: {border};
                    border-radius: 11px;
                    font-size: 11px;
                    color: #1c1c1e;
                }}
                QPushButton:hover {{ border: 2px solid rgba(0,0,0,0.36); }}
            """)

    def _apply_bg(self, color: NoteColor):
        hex_color = COLOR_HEX[color]
        self._body_widget.setStyleSheet(f"background: {hex_color};")
        self._footer_widget.setStyleSheet(
            f"background: {hex_color}; border-top: 1px solid #d1d1d6;"
        )

    def _on_save(self):
        self._do_save()
        self.back_requested.emit()

    def _on_delete(self):
        msg = QMessageBox(self)
        msg.setWindowTitle(_("Delete Note"))
        msg.setText(_("Are you sure you want to delete this note?\nThis action cannot be undone."))
        msg.setStandardButtons(
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        msg.setDefaultButton(QMessageBox.StandardButton.No)
        msg.button(QMessageBox.StandardButton.Yes).setText(_("Delete"))
        msg.button(QMessageBox.StandardButton.No).setText(_("Cancel"))
        if msg.exec() != QMessageBox.StandardButton.Yes:
            return
        if self._note_id:
            self._client.delete_note(self._note_id)
        self.back_requested.emit()

    def _on_delete_image(self):
        if self._note_id and self._image_url:
            self._client.delete_image(self._note_id, self._image_url)
            self._image_url = None
            self._img_section.hide()

    def _on_copy_to_clipboard(self):
        title, body = self._note_content()
        parts = [p for p in [title, body] if p]
        QApplication.clipboard().setText("\n\n".join(parts))
        self._show_export_toast(f"✓  {_('Copied to clipboard')}")

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape:
            self._on_back()
        super().keyPressEvent(event)

    # ── 내보내기 / 공유 ───────────────────────────────────────

    def _note_content(self) -> tuple[str, str]:
        """(title, body_text) 반환. 체크리스트는 텍스트로 직렬화."""
        title = self._title_edit.text().strip()
        if self._list_scroll.isVisible():
            body = "\n".join(
                f"{'[x]' if cb.isChecked() else '[ ]'} {text}"
                for cb, text in self._checklist_rows
            )
        else:
            body = self._body_edit.toPlainText().strip()
        return title, body

    def _on_export_click(self):
        menu = QMenu(self)
        menu.setStyleSheet("""
            QMenu {
                background: white; border: 1px solid #d1d1d6;
                border-radius: 8px; padding: 4px; font-size: 12px;
            }
            QMenu::item { padding: 6px 16px; border-radius: 4px; }
            QMenu::item:selected { background: #f2f2f7; color: #1c1c1e; }
            QMenu::separator { height: 1px; background: #d1d1d6; margin: 4px 8px; }
            QMenu::right-arrow { image: none; width: 8px; }
        """)

        save_menu = menu.addMenu(f"💾  {_('Save as File')}")
        save_menu.setStyleSheet(menu.styleSheet())
        save_menu.addAction("Markdown (.md)").triggered.connect(self._on_save_md)
        save_menu.addAction("Text (.txt)").triggered.connect(self._on_save_txt)
        save_menu.addAction("Image (.png)").triggered.connect(self._on_save_png)
        save_menu.addAction("PDF (.pdf)").triggered.connect(self._on_save_pdf)
        save_menu.addAction("한글 (.hwpx)").triggered.connect(self._on_save_hwpx)
        save_menu.addAction("Word (.docx)").triggered.connect(self._on_save_docx)

        email_act = menu.addAction(f"✉  {_('Send via Email')}")
        email_act.triggered.connect(self._on_email_share)

        my_email_act = menu.addAction(f"📨  {_('Send to My Email')}")
        my_email_act.triggered.connect(self._on_my_email_share)

        kakao_act = menu.addAction(f"💬  {_('Share via KakaoTalk')}")
        kakao_act.triggered.connect(self._on_kakao_share)

        menu.addSeparator()

        translate_menu = menu.addMenu(f"🌐  {_('Translate & New Note')}")
        translate_menu.setStyleSheet(menu.styleSheet())
        current_lang = i18n.current_lang()
        for code, name in i18n.SUPPORTED_LANGS.items():
            if code != current_lang:
                act = translate_menu.addAction(f"→ {name}")
                act.triggered.connect(
                    lambda checked=False, c=code: self._on_translate(c)
                )
        if self._is_new:
            translate_menu.setEnabled(False)

        pos = self._export_btn.mapToGlobal(self._export_btn.rect().topLeft())
        menu.exec(pos)

    # ── 파일 저장 공통 헬퍼 ─────────────────────────────────

    def _fetch_note_image(self) -> bytes | None:
        if not self._image_url:
            return None
        return self._client.fetch_image(self._image_url)

    def _note_html(self, title: str, body: str, img_data: bytes | None = None) -> str:
        """PNG/PDF 렌더링용 HTML 생성."""
        bg = COLOR_HEX.get(self._current_color, "#FFFFFF")
        parts = [f'<html><body style="font-family:-apple-system,sans-serif;'
                 f'background:{bg};padding:24px;max-width:600px;">']
        if title:
            parts.append(f'<h1 style="font-size:18px;margin:0 0 12px 0">{_he(title)}</h1>')
        if img_data:
            b64 = base64.b64encode(img_data).decode()
            parts.append(f'<img src="data:image/png;base64,{b64}"'
                         f' style="max-width:100%;margin-bottom:12px"><br>')
        if self._list_scroll.isVisible():
            parts.append('<div style="line-height:1.8">')
            for cb, text in self._checklist_rows:
                sym = "☑" if cb.isChecked() else "☐"
                parts.append(f'{sym} {_he(text)}<br>')
            parts.append('</div>')
        elif body:
            safe = _he(body).replace("\n", "<br>")
            parts.append(f'<p style="white-space:pre-wrap;line-height:1.6;margin:0">{safe}</p>')
        parts.append('</body></html>')
        return "".join(parts)

    # ── #49/#50 파일로 저장 ───────────────────────────────────

    def _on_save_md(self):
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.md")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "Markdown (*.md);;All files (*)"
        )
        if not path:
            return
        p = Path(path)
        lines: list[str] = []
        if title:
            lines += [f"# {title}", ""]
        img_data = self._fetch_note_image()
        if img_data:
            img_dir = p.parent / f"{p.stem}_images"
            img_dir.mkdir(exist_ok=True)
            (img_dir / "image.png").write_bytes(img_data)
            lines += [f"![image]({p.stem}_images/image.png)", ""]
        if self._list_scroll.isVisible():
            for cb, text in self._checklist_rows:
                lines.append(f"- [{'x' if cb.isChecked() else ' '}] {text}")
        elif body:
            lines.append(body)
        p.write_text("\n".join(lines), encoding="utf-8")
        self._show_export_toast(f"✓  {p.name}")

    def _on_save_txt(self):
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.txt")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "Text (*.txt);;All files (*)"
        )
        if not path:
            return
        lines: list[str] = []
        if title:
            lines += [title, "─" * min(len(title) * 2, 40), ""]
        if self._list_scroll.isVisible():
            for cb, text in self._checklist_rows:
                lines.append(f"{'☑' if cb.isChecked() else '☐'} {text}")
        elif body:
            lines.append(body)
        if self._image_url:
            lines += ["", "[이미지 첨부]"]
        Path(path).write_text("\n".join(lines), encoding="utf-8")
        self._show_export_toast(f"✓  {Path(path).name}")

    def _on_save_png(self):
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.png")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "Image (*.png);;All files (*)"
        )
        if not path:
            return
        img_data = self._fetch_note_image()
        html = self._note_html(title, body, img_data)
        doc = QTextDocument()
        doc.setHtml(html)
        doc.setPageSize(QSizeF(660, 99999))
        doc.setPageSize(QSizeF(660, doc.size().height()))
        W = 660; H = int(doc.size().height()) + 48
        img = QImage(W, H, QImage.Format.Format_ARGB32)
        img.fill(QColor(COLOR_HEX.get(self._current_color, "#FFFFFF")))
        painter = QPainter(img)
        painter.translate(0, 24)
        doc.drawContents(painter)
        painter.end()
        img.save(path, "PNG")
        self._show_export_toast(f"✓  {Path(path).name}")

    def _on_save_pdf(self):
        from PyQt6.QtPrintSupport import QPrinter
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.pdf")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "PDF (*.pdf);;All files (*)"
        )
        if not path:
            return
        img_data = self._fetch_note_image()
        html = self._note_html(title, body, img_data)
        printer = QPrinter()
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(path)
        doc = QTextDocument()
        doc.setHtml(html)
        doc.print(printer)
        self._show_export_toast(f"✓  {Path(path).name}")

    def _on_save_hwpx(self):
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.hwpx")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "한글 (*.hwpx);;All files (*)"
        )
        if not path:
            return
        checklist = None
        if self._list_scroll.isVisible():
            checklist = [(cb.isChecked(), text) for cb, text in self._checklist_rows]
        img_data = self._fetch_note_image()
        _write_hwpx(path, title, body, checklist, img_data)
        self._show_export_toast(f"✓  {Path(path).name}")

    def _on_save_docx(self):
        title, body = self._note_content()
        stem = title or f"note_{(self._note_id or 'new')[:8]}"
        default = str(Path.home() / "Downloads" / f"{stem}.docx")
        path, _filter = QFileDialog.getSaveFileName(
            self, _("Save as File"), default, "Word (*.docx);;All files (*)"
        )
        if not path:
            return
        checklist = None
        if self._list_scroll.isVisible():
            checklist = [(cb.isChecked(), text) for cb, text in self._checklist_rows]
        img_data = self._fetch_note_image()
        _write_docx(path, title, body, checklist, img_data)
        self._show_export_toast(f"✓  {Path(path).name}")

    # ── #27 이메일 공유 ──────────────────────────────────────

    def _on_email_share(self):
        from keep4mac.core import settings as _settings
        title, body = self._note_content()
        subject = f"keep4mac - {title}" if title else "keep4mac"
        url = QUrl(f"mailto:?subject={quote(subject)}&body={quote(body)}")
        QDesktopServices.openUrl(url)

    def _on_my_email_share(self):
        from keep4mac.core import settings as _settings
        to = _settings.get_my_email()
        if not to:
            self._show_export_toast(f"⚠  {_('Set your email in Settings first')}", duration=3000)
            return
        title, body = self._note_content()
        subject = f"keep4mac - {title}" if title else "keep4mac"
        url = QUrl(f"mailto:{quote(to)}?subject={quote(subject)}&body={quote(body)}")
        QDesktopServices.openUrl(url)

    # ── #26 카카오톡 공유 ────────────────────────────────────

    def _on_kakao_share(self):
        title, body = self._note_content()
        text = f"{title}\n\n{body}".strip() if title else body
        QApplication.clipboard().setText(text)
        QDesktopServices.openUrl(QUrl("kakaotalk://"))
        self._show_export_toast(_("Copied to clipboard. Open KakaoTalk to share."))

    # ── #32 번역 새 노트 ─────────────────────────────────────

    def _on_translate(self, target_lang: str):
        title, content = self._note_content()
        source = i18n.current_lang()
        self._export_btn.setEnabled(False)
        self._show_export_toast(f"⏳  {_('Translating…')}")
        self._translate_thread = _TranslateThread(title, content, source, target_lang)
        self._translate_thread.done.connect(
            lambda t, c: self._on_translate_done(t, c, target_lang)
        )
        self._translate_thread.error.connect(
            lambda msg: self._on_translate_error(msg)
        )
        self._translate_thread.start()

    def _on_translate_error(self, msg: str):
        self._export_btn.setEnabled(True)
        self._show_export_toast(f"⚠  {_('Translation failed')}: {msg[:60]}", duration=5000)

    def _on_translate_done(self, t_title: str, t_content: str, lang_code: str):
        prefix = f"[{lang_code.upper()}] "
        title = prefix + t_title if t_title else prefix.strip()
        self._show_export_toast(f"📝  {_('Creating note…')}")
        self._create_note_thread = _CreateNoteThread(
            self._client, title, t_content, self._current_color
        )
        self._create_note_thread.done.connect(self.back_requested.emit)
        self._create_note_thread.error.connect(
            lambda msg: self._on_create_note_error(msg)
        )
        self._create_note_thread.finished.connect(
            lambda: self._export_btn.setEnabled(True)
        )
        self._create_note_thread.start()

    def _on_create_note_error(self, msg: str):
        self._show_export_toast(f"⚠  {_('Failed to create note')}: {msg[:60]}", duration=5000)

    # ── 내부 토스트 ──────────────────────────────────────────

    def _show_export_toast(self, message: str, duration: int = 2500):
        if hasattr(self, "_toast") and self._toast is not None:
            try:
                self._toast.deleteLater()
            except RuntimeError:
                pass
        toast = QLabel(message, self)
        self._toast = toast
        toast.setAlignment(Qt.AlignmentFlag.AlignCenter)
        toast.setStyleSheet("""
            QLabel {
                background: rgba(32,33,36,0.88); color: white;
                font-size: 12px; border-radius: 8px; padding: 8px 14px;
            }
        """)
        toast.adjustSize()
        w = min(toast.sizeHint().width() + 24, self.width() - 32)
        toast.setFixedWidth(w)
        toast.move((self.width() - w) // 2, self.height() - toast.sizeHint().height() - 20)
        toast.show()
        toast.raise_()
        QTimer.singleShot(duration, toast.deleteLater)
